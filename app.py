import os
import re
import json
import time
import uuid
import threading
import glob
import zipfile
from queue import Queue
from datetime import datetime

from flask import Flask, request, jsonify, render_template, Response, send_file
from werkzeug.utils import secure_filename
import pandas as pd
import openpyxl
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
import docx

# ==========================================
# KONFIGURÁCIÓ
# ==========================================
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['DOWNLOAD_FOLDER'] = 'downloads'
app.config['PROMPTS_FOLDER'] = 'prompts'
app.config['JOBS_FOLDER'] = 'jobs'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB limit

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['DOWNLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['PROMPTS_FOLDER'], exist_ok=True)
os.makedirs(app.config['JOBS_FOLDER'], exist_ok=True)

@app.after_request
def add_header(response):
    response.headers['X-Frame-Options'] = 'ALLOWALL'
    response.headers['Access-Control-Allow-Origin'] = '*'
    return response

try:
    client = OpenAI()
except Exception as e:
    print(f"Figyelem: OpenAI kliens inicializálási hiba: {e}")
    client = None

AVAILABLE_MODELS = [
    "gpt-5.4",
    "gpt-5.4-mini",
    "gpt-5.4-nano",
    "gpt-5.4-pro",
    "gpt-4.1",
    "gpt-4.1-mini",
    "gpt-4.1-nano",
]
DEFAULT_MODEL = "gpt-5.4-mini"

FORBIDDEN_PHRASES = [
    "Az oldal szerint", "A leírás szerint", "Ez nem bénázás",
    "Tapasztalatból mondom", "ez nem csak", "talán", "nagyjából", "bizonyos értelemben"
]

# A memóriában tartott job esemény sorok az SSE-hez
job_events = {}
# Lock a párhuzamos generálás megakadályozására
generation_lock = threading.Lock()

# ==========================================
# JOB KEZELÉS (PERZISZTENS)
# ==========================================
def get_job_path(job_id):
    return os.path.join(app.config['JOBS_FOLDER'], f"{job_id}.json")

def save_job(job_data):
    job_id = job_data['job_id']
    with open(get_job_path(job_id), 'w', encoding='utf-8') as f:
        json.dump(job_data, f, ensure_ascii=False, indent=2)

def load_job(job_id):
    path = get_job_path(job_id)
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return None

def get_all_jobs():
    jobs = []
    for path in glob.glob(os.path.join(app.config['JOBS_FOLDER'], "*.json")):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                jobs.append(json.load(f))
        except:
            pass
    # Rendezzük csökkenő sorrendbe (legújabb elöl)
    jobs.sort(key=lambda x: x.get('started_at', ''), reverse=True)
    return jobs

def emit_event(job_id, event_data):
    if job_id in job_events:
        job_events[job_id].put(event_data)

# ==========================================
# TONE GUIDE ÉS PIPELINE FÁJLRENDSZER
# ==========================================
TONE_GUIDE_PATH = os.path.join(app.config['PROMPTS_FOLDER'], 'tone_guide.txt')
PIPELINE_PATH = os.path.join(app.config['PROMPTS_FOLDER'], 'pipeline.json')

DEFAULT_TONE_GUIDE = """1. Magas szintű áttekintés
Általános stílusösszefoglaló
Ez a stílus közérthető, pragmatikus, döntést segítő szakmai tartalomírás. Nem irodalmi, nem játékos, és nem is akadémikus: inkább olyan, mintha egy tapasztalt tanácsadó leülne melléd, és gyorsan rendet rakna a témában. A mondanivaló mindig lebontott, „megfogható", kézzelfogható szempontokra van osztva, ezért az olvasó ritkán marad absztrakt állításokkal egyedül. A forma végig azt sugallja: „itt most tisztázzuk a helyzetet, és kapsz egy biztonságos következő lépést."

Szerzői persona (következtetve)
A domináns persona: szakértő tanácsadó + türelmes magyarázó + finoman értékesítő konzultáns. A szerző nem haverkodik, és nem fölülről beszél: inkább „kompetens vezetőként" vagy „józan eligazítóként" szólal meg.

Az írás elsődleges céljai
Az elsődleges célok: tisztázni, bizonytalanságot csökkenteni, szakértői hitelességet építeni, majd óvatosan konverzióba terelni. A legtöbb cikk előbb oktat, aztán keretez, és csak utána ajánl szolgáltatást vagy következő lépést. Ezért a meggyőzés itt nem agresszív, hanem strukturális: a józan rendrakásból nő ki.

2. Hang és tónus
Hang
A hang félformális, közvetlen, társalgóan szakmai. Nem laza, de nem merev. Gyakori a második személyű megszólítás ("neked", "nálad", "ha szeretnéd"), illetve a vállalati többes első személy ("bemutatjuk", "segítünk", "átnézzük"), ami egyszerre sugall szakértelmet és szolgáltatói jelenlétet.

8. Stílusszabályok / Ellenőrzőlista
Hangszabályok
- Beszélj úgy, mint egy hozzáértő tanácsadó, ne úgy, mint egy reklámszövegíró, aki túl korán eladni akar.
- Az első 2–4 mondatban nevezd meg az olvasó valós helyzetét vagy félreértését.
- Használj közvetlen, de nem haverkodó megszólalást.
- Mutass empátiát a problémára, de maradj rendezett és tárgyszerű.
- Legyen a hangod magabiztos, de ne hangoskodó.
- Inkább eligazíts, mint lenyűgözz.
- A hitelességet konkrétumokkal építsd, ne önfényezéssel.
- A CTA előtt adj valódi értéket és egyértelmű gondolati rendet.

NE tedd
- Ne írj hosszú, díszes, többszörösen alárendelt mondatokat.
- Ne használj túl sok bizonytalanító puhaságot, mint a „talán", „nagyjából", „bizonyos értelemben".
- Ne menj át teljesen sales-es hype-ba az első bekezdésekben.
- Ne hagyd, hogy egy szakasz cím nélkül, tömbszövegként ömöljön rá az olvasóra.
- Ne legyen ilyen stílusú mondat: „ez nem csak X, hanem Y"
- Felkiáltójel szinte ne legyen."""

DEFAULT_PIPELINE = {
    "steps": [
        {
            "id": 1,
            "name": "Cikk generálás",
            "enabled": True,
            "prompt": """Te egy senior SEO-szövegíró és tartalomstratéga vagy. Magyarul írsz, természetesen, marketing-szag nélkül, de konverzióra optimalizálva.
Írj egy minőségi, keresőoptimalizált blogcikket a következő szolgáltatás oldal támogatására, és organikus forgalom + érdeklődők szerzésére.

Cég honlapjának url-je: {ceg_url}
Cikk címe: {cikk_cim}

Helyezz el rajta linkeket az alábbi kulcsszavakról az adott oldalakra, összesen {link_db} link legyen egy cikkben, a linkbe semmiképp ne illessz utm paramétert (Markdown formátumban illeszd be a linkeket a szövegbe):
{linkek_felsorolasa}

{opcionalis_korabbi_cikkek_blokk}

Követelmények a cikkhez
Nyelv: magyar
Terjedelem: 600-1000 szó
E-E-A-T: jelezd a szakértelmet (folyamat, módszer, tapasztalati jelek), kerüld a túlzó ígéreteket.
Kimenet:
Csak a teljes cikket add vissza.
Ne adj title, meta, outline elemzést.
Az alcímeket Markdown formátumban add meg: ## Alcím (két hashtaggel), a szövegben ne szerepeljen a főcím (azt külön adjuk hozzá).

Fontos:
Ne tömj kulcsszót, legyen természetes.
Ne találj ki konkrét statisztikát; ha számot írsz, legyen általános vagy jelöld természetes módon, hogy "példa".
Ha írsz állításokat a céggel kapcsolatban, csakis valósakat írj, olyanokat, melyek megtalálhatók a cég honlapján, ne találj ki nem valós információkat.
Ne legyenek benne ilyen jellegű megfogalmazások:
"Az oldal szerint …", "A leírás szerint …"
"Ez nem bénázás, hanem profizmus."
"Tapasztalatból mondom"
Ne tegyél be forrás hivatkozásokat a bekezdések végére, csak az általam kért hivatkozások szerepeljenek a szövegben.
Ne legyél túlságosan közvetlen.
Ne használj ilyen szerkezetű mondatot: "ez nem csak X, hanem Y".
Ne használj bizonytalanító szavakat döntéstámogató szövegben: talán, nagyjából, bizonyos értelemben.

Stílus és hangnem útmutató:
{tone_guide}{megjegyzes_blokk}"""
        },
        {
            "id": 2,
            "name": "Tényellenőrzés és javítás",
            "enabled": True,
            "prompt": """Az alábbi cikket ellenőrizd le két szempontból:

1. TÉNYELLENŐRZÉS: Tartalmaz-e olyan konkrét állítást a cégről ({ceg_url}), amely nyilvánvalóan kitalált vagy nem ellenőrizhető? Csak egyértelmű kitalált állítást jelezz.

2. FORMÁTUM ÉS NYELVEZET: Szerepelnek-e tiltott fordulatok? Tiltott: "Az oldal szerint", "A leírás szerint", "Ez nem bénázás", "Tapasztalatból mondom", "ez nem csak", "talán", "nagyjából", "bizonyos értelemben". Van-e legalább 2 db ## jelű H2 alcím?

Ha mindkét szempont rendben van, add vissza a cikket változatlanul.
Ha valamelyik szempontban probléma van, javítsd ki a problémát és add vissza a teljes javított cikket.

Fontos: CSAK a cikk szövegét add vissza, semmi mást (sem magyarázatot, sem megjegyzést).

Cikk:
{aktualis_cikk}"""
        },
        {
            "id": 3,
            "name": "Link és UTM ellenőrzés és javítás",
            "enabled": True,
            "prompt": """Az alábbi cikket ellenőrizd le és szükség esetén javítsd:

1. UTM PARAMÉTEREK: Ha bármely linkben utm_ paraméter szerepel, távolítsd el.
2. ELVÁRT LINKEK: Szerepelnek-e a cikkben a következő kulcsszavak/linkek: {elvart_linkek}? Ha valamelyik hiányzik, illeszd be természetes módon a szövegbe.

Add vissza a teljes (javított vagy változatlan) cikket. CSAK a cikk szövegét add vissza, semmi mást.

Cikk:
{aktualis_cikk}"""
        }
    ],
    "versions": []
}

def load_tone_guide():
    if os.path.exists(TONE_GUIDE_PATH):
        with open(TONE_GUIDE_PATH, 'r', encoding='utf-8') as f:
            return f.read()
    return DEFAULT_TONE_GUIDE

def save_tone_guide(text):
    with open(TONE_GUIDE_PATH, 'w', encoding='utf-8') as f:
        f.write(text)

def load_pipeline_data():
    if os.path.exists(PIPELINE_PATH):
        with open(PIPELINE_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    return DEFAULT_PIPELINE.copy()

def save_pipeline_data(steps):
    data = load_pipeline_data()
    old_steps = data.get('steps', [])
    versions = data.get('versions', [])

    if old_steps:
        versions.append({
            'version': len(versions) + 1,
            'steps': old_steps,
            'saved_at': datetime.now().isoformat(timespec='seconds')
        })

    data['steps'] = steps
    data['versions'] = versions

    with open(PIPELINE_PATH, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def init_files():
    if not os.path.exists(TONE_GUIDE_PATH):
        save_tone_guide(DEFAULT_TONE_GUIDE)
    if not os.path.exists(PIPELINE_PATH):
        with open(PIPELINE_PATH, 'w', encoding='utf-8') as f:
            json.dump(DEFAULT_PIPELINE, f, ensure_ascii=False, indent=2)

init_files()

# ==========================================
# SEGÉDFÜGGVÉNYEK (Word formázás)
# ==========================================
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    color = docx.oxml.shared.OxmlElement('w:color')
    color.set(docx.oxml.shared.qn('w:val'), '0000FF')
    rPr.append(color)

    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    text_node = docx.oxml.shared.OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def add_formatted_runs(p, line):
    pattern = r'(\*\*.*?\*\*|\[.*?\]\(.*?\))'
    parts = re.split(pattern, line)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('[') and '](' in part and part.endswith(')'):
            text_end = part.find('](')
            link_text = part[1:text_end]
            url = part[text_end+2:-1]
            add_hyperlink(p, url, link_text)
        else:
            p.add_run(part)

def format_markdown_to_docx(doc, text):
    lines = text.split('\n')

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith('### '):
            heading = doc.add_heading(stripped[4:], level=3)
            heading.style.font.color.rgb = RGBColor(0, 0, 0)
        elif stripped.startswith('## '):
            heading = doc.add_heading(stripped[3:], level=2)
            heading.style.font.color.rgb = RGBColor(0, 0, 0)
        elif stripped.startswith('# '):
            heading = doc.add_heading(stripped[2:], level=1)
            heading.style.font.color.rgb = RGBColor(0, 0, 0)
        elif re.match(r'^[-*] ', stripped):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_runs(p, stripped[2:])
        elif re.match(r'^\d+\.\s\*\*', stripped):
            content = re.sub(r'^\d+\.\s', '', stripped)
            heading_text = content.strip('*').strip()
            heading = doc.add_heading(heading_text, level=3)
            heading.style.font.color.rgb = RGBColor(0, 0, 0)
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            content = re.sub(r'^\d+\.\s', '', stripped)
            add_formatted_runs(p, content)
        elif re.match(r'^\*\*\d+\.', stripped):
            heading_text = stripped.strip('*').strip()
            heading = doc.add_heading(heading_text, level=3)
            heading.style.font.color.rgb = RGBColor(0, 0, 0)
        else:
            p = doc.add_paragraph()
            add_formatted_runs(p, stripped)

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

# ==========================================
# CIKK GENERÁLÓ LOGIKA
# ==========================================
def safe_format(template, variables):
    """Biztonságos string formázás, ami figyelmen kívül hagyja a hiányzó változókat."""
    class SafeDict(dict):
        def __missing__(self, key):
            return '{' + key + '}'
    return template.format_map(SafeDict(variables))

def call_llm(prompt_text, model, temperature=0.7, retries=3):
    """LLM hívás rate limit és timeout kezeléssel."""
    for attempt in range(retries):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt_text}],
                temperature=temperature,
                timeout=60.0
            )
            content = response.choices[0].message.content.strip()
            if not content and attempt < retries - 1:
                print(f"Üres válasz kapva, újrapróbálkozás ({attempt+1}/{retries})...")
                time.sleep(2)
                continue
            return content
        except Exception as e:
            error_str = str(e)
            if "429" in error_str or "rate limit" in error_str.lower():
                wait_time = 30 * (attempt + 1)
                print(f"Rate limit elérve. Várakozás {wait_time} másodpercet ({attempt+1}/{retries})...")
                time.sleep(wait_time)
            elif "timeout" in error_str.lower():
                print(f"Timeout hiba. Újrapróbálkozás ({attempt+1}/{retries})...")
                time.sleep(5)
            else:
                if attempt == retries - 1:
                    return f"API Hiba: {error_str}"
                time.sleep(2)
    return "API Hiba: Többszöri próbálkozás után sem sikerült választ kapni."

def generate_single_article(row_data, job_id, row_index, model):
    job = load_job(job_id)
    if not job:
        return None

    def update_status(status, message):
        # Update in-memory job state and save to file
        job = load_job(job_id)
        job['rows'][row_index]['status'] = status
        if message:
            job['rows'][row_index]['message'] = message
        
        # Add to steps_log
        if 'steps_log' not in job['rows'][row_index]:
            job['rows'][row_index]['steps_log'] = []
        job['rows'][row_index]['steps_log'].append(f"{datetime.now().strftime('%H:%M:%S')} - {status}: {message}")
        
        save_job(job)
        
        # Emit event to frontend
        emit_event(job_id, {
            'type': 'row_update',
            'row_index': row_index,
            'status': status,
            'message': message
        })

    # Ellenőrzés: kötelező mezők
    ceg_url = row_data.get('ceg_url', '').strip()
    cikk_cim = row_data.get('cikk_cim', '').strip()
    
    if not ceg_url or not cikk_cim:
        update_status('Hiba', 'Hiányzó cég URL vagy cikk cím')
        job = load_job(job_id)
        job['rows'][row_index]['status'] = 'error'
        save_job(job)
        return None

    # Változók előkészítése
    vars_dict = {k: str(v) for k, v in row_data.items()}

    links = []
    keywords = []
    for i in range(1, 6):
        kw = row_data.get(f'link_{i}_kulcsszo', '')
        url = row_data.get(f'link_{i}_url', '')
        if kw and url:
            links.append(f"- [{kw}]({url})")
            keywords.append(kw)

    vars_dict['link_db'] = str(len(links))
    vars_dict['linkek_felsorolasa'] = "\n".join(links)
    vars_dict['elvart_linkek'] = ", ".join(keywords) if keywords else "nincs megadva"

    korabbi_cikkek = []
    for i in range(1, 3):
        url = row_data.get(f'korabbi_cikk_url_{i}', '')
        if url:
            korabbi_cikkek.append(f"- {url}")

    if korabbi_cikkek:
        vars_dict['opcionalis_korabbi_cikkek_blokk'] = (
            "Kérlek, helyezz el természetes hivatkozást a következő korábbi cikk(ek)re is "
            "(a megfelelő kontextusban):\n" + "\n".join(korabbi_cikkek)
        )
    else:
        vars_dict['opcionalis_korabbi_cikkek_blokk'] = ""

    megjegyzes = row_data.get('megjegyzes', '')
    vars_dict['megjegyzes_blokk'] = f"\nEgyéb instrukciók ehhez a cikkhez:\n{megjegyzes}" if megjegyzes else ""
    vars_dict['tone_guide'] = load_tone_guide()
    vars_dict['aktualis_cikk'] = ""
    vars_dict['elozo_kimenet'] = ""

    # Pipeline betöltése
    pipeline = load_pipeline_data()
    steps = [s for s in pipeline.get('steps', []) if s.get('enabled', True)]

    if not steps:
        update_status('Hiba', 'Nincs engedélyezett lépés a pipeline-ban!')
        job = load_job(job_id)
        job['rows'][row_index]['status'] = 'error'
        save_job(job)
        return None

    update_status('Folyamatban', 'Pipeline indítása...')

    # Szekvenciális futtatás – minden lépés sorban fut
    for step_num, step in enumerate(steps, start=1):
        step_name = step.get('name', f'{step_num}. lépés')
        update_status('Folyamatban', f'{step_name}...')

        prompt = safe_format(step['prompt'], vars_dict)
        output = call_llm(prompt, model, temperature=0.7)

        if output.startswith("API Hiba:"):
            update_status('Hiba', f'{step_name}: {output}')
            job = load_job(job_id)
            job['rows'][row_index]['status'] = 'error'
            save_job(job)
            return None

        # Változók frissítése a következő lépésekhez
        step_id = step.get('id', step_num)
        vars_dict[f'lepes_{step_id}_kimenet'] = output
        vars_dict[f'lepes_{step_num}_kimenet'] = output  # sorszám alapján is
        vars_dict['elozo_kimenet'] = output
        vars_dict['aktualis_cikk'] = output

        time.sleep(1)

    # Az utolsó lépés kimenete lesz a végső cikk
    final_article = vars_dict['aktualis_cikk']

    # Kód alapú utóellenőrzés (UTM, szószám)
    warnings = []
    if "utm_" in final_article.lower():
        warnings.append("Figyelem: utm_ paraméter maradt a cikkben")

    word_count = len(re.findall(r'\b\w+\b', final_article))
    if word_count < 600:
        warnings.append(f"Rövid cikk ({word_count} szó)")
    elif word_count > 1000:
        warnings.append(f"Hosszú cikk ({word_count} szó)")

    if warnings:
        update_status('Kész', 'Kész (figyelmeztetések: ' + '; '.join(warnings) + ')')
    else:
        update_status('Kész', f'Sikeres generálás ({word_count} szó)')

    # Véglegesítés a job fájlban
    job = load_job(job_id)
    job['rows'][row_index]['status'] = 'done'
    job['rows'][row_index]['article'] = final_article
    job['rows'][row_index]['completed_at'] = datetime.now().isoformat()
    save_job(job)

    return final_article

def create_output_files(job_id):
    job = load_job(job_id)
    if not job:
        return
        
    doc = Document()
    sikeres_cikkek = 0
    hibas_cikkek = 0
    
    # Készítsünk egy txt fallback zipet is
    zip_filename = f"cikkek_{job_id}.zip"
    zip_filepath = os.path.join(app.config['DOWNLOAD_FOLDER'], zip_filename)
    
    with zipfile.ZipFile(zip_filepath, 'w') as zipf:
        for row in job['rows']:
            if row.get('status') == 'done' and row.get('article'):
                sikeres_cikkek += 1
                cikk_cim = row.get('cikk_cim', f'cikk_{row.get("index", sikeres_cikkek)}')
                safe_cim = sanitize_filename(cikk_cim)
                
                # Word doc építése
                if sikeres_cikkek > 1:
                    doc.add_page_break()
                
                heading = doc.add_heading(cikk_cim, 0)
                heading.style.font.color.rgb = RGBColor(0, 0, 0)
                format_markdown_to_docx(doc, row['article'])
                
                # Txt zip építése
                zipf.writestr(f"{safe_cim}.txt", row['article'])
            else:
                hibas_cikkek += 1
                
    # Összefoglaló oldal beszúrása a Word elejére (ha van hiba)
    if hibas_cikkek > 0 and sikeres_cikkek > 0:
        # A docx python library nem támogatja egyszerűen az oldal beszúrását az elejére, 
        # ezért egy új dokumentumot hozunk létre és abba másoljuk.
        # Egyszerűbb megoldás: a végére tesszük az összefoglalót.
        doc.add_page_break()
        doc.add_heading('Generálási összefoglaló', level=1)
        doc.add_paragraph(f"Sikeresen generált cikkek: {sikeres_cikkek}")
        doc.add_paragraph(f"Hibával végződött sorok: {hibas_cikkek}")
        
        p = doc.add_paragraph()
        p.add_run("A hibás sorok részletei a webes felületen tekinthetők meg.")

    if sikeres_cikkek > 0:
        docx_filename = f"keszult_cikkek_{job_id}.docx"
        docx_filepath = os.path.join(app.config['DOWNLOAD_FOLDER'], docx_filename)
        doc.save(docx_filepath)
        
        job['download_url'] = f"/download/{docx_filename}"
        job['zip_url'] = f"/download/{zip_filename}"
        save_job(job)

def generation_worker(job_id):
    # Lockoljuk a futást
    acquired = generation_lock.acquire(blocking=False)
    if not acquired:
        job = load_job(job_id)
        job['status'] = 'error'
        job['error_message'] = 'Már fut egy generálás. Kérlek várj.'
        save_job(job)
        emit_event(job_id, {'type': 'error', 'message': 'Már fut egy generálás.'})
        return
        
    try:
        job = load_job(job_id)
        if not job:
            return

        job['status'] = 'running'
        save_job(job)
        
        rows = job['rows']
        model = job.get('model', DEFAULT_MODEL)

        for idx, row in enumerate(rows):
            # Csak a nem befejezett sorokat dolgozzuk fel
            if row.get('status') == 'done':
                continue
                
            # Kihagyjuk az üres sorokat
            if not row.get('ceg_url') and not row.get('cikk_cim'):
                job = load_job(job_id)
                job['rows'][idx]['status'] = 'done'
                job['rows'][idx]['message'] = 'Üres sor kihagyva'
                save_job(job)
                
                job['completed_rows'] += 1
                save_job(job)
                emit_event(job_id, {
                    'type': 'progress',
                    'completed': job['completed_rows'],
                    'total': job['total_rows']
                })
                continue

            # Állapot frissítés futóra
            job = load_job(job_id)
            job['rows'][idx]['status'] = 'running'
            save_job(job)
            
            generate_single_article(row, job_id, idx, model)

            # Újratöltjük a jobot, hogy megnézzük a végső státuszt
            job = load_job(job_id)
            job['completed_rows'] += 1
            save_job(job)
            
            emit_event(job_id, {
                'type': 'progress',
                'completed': job['completed_rows'],
                'total': job['total_rows']
            })

        # Fájlok generálása
        create_output_files(job_id)

        job = load_job(job_id)
        job['status'] = 'done'
        save_job(job)
        
        emit_event(job_id, {
            'type': 'complete',
            'download_url': job.get('download_url', None),
            'zip_url': job.get('zip_url', None)
        })
        emit_event(job_id, None)  # Stream vége
        
    except Exception as e:
        print(f"Hiba a workerben: {e}")
        job = load_job(job_id)
        if job:
            job['status'] = 'error'
            job['error_message'] = str(e)
            save_job(job)
            emit_event(job_id, {'type': 'error', 'message': f'Rendszerhiba: {str(e)}'})
            emit_event(job_id, None)
    finally:
        generation_lock.release()

# ==========================================
# ÚTVONALAK – ALAP
# ==========================================
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'Nincs fájl kiválasztva'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nincs fájl kiválasztva'}), 400

    if not file.filename.endswith('.xlsx'):
        return jsonify({'error': 'Csak .xlsx fájl tölthető fel'}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    try:
        df = pd.read_excel(filepath)
        df = df.fillna('')
        
        # Ellenőrizzük a kötelező oszlopokat
        required_cols = ['ceg_url', 'cikk_cim']
        missing_cols = [col for col in required_cols if col not in df.columns]
        if missing_cols:
            return jsonify({'error': f'Hiányzó kötelező oszlopok: {", ".join(missing_cols)}'}), 400

        rows = []
        for idx, row in df.iterrows():
            row_dict = row.to_dict()
            row_dict['index'] = idx
            row_dict['status'] = 'pending'
            row_dict['message'] = 'Várakozik'
            rows.append(row_dict)

        columns = [col for col in df.columns]

        return jsonify({
            'success': True,
            'columns': columns,
            'rows': rows
        })
    except Exception as e:
        return jsonify({'error': f'Hiba a fájl feldolgozásakor: {str(e)}'}), 500

@app.route('/start-generation', methods=['POST'])
def start_generation():
    if generation_lock.locked():
        return jsonify({'error': 'Már fut egy generálás. Kérlek várj, amíg befejeződik.'}), 409
        
    data = request.json
    rows = data.get('rows', [])
    model = data.get('model', DEFAULT_MODEL)

    if model not in AVAILABLE_MODELS:
        model = DEFAULT_MODEL

    if not rows:
        return jsonify({'error': 'Nincsenek adatok a generáláshoz'}), 400

    job_id = str(uuid.uuid4())
    
    # Töröljük az üres sorokat a végéről
    while rows and not rows[-1].get('ceg_url') and not rows[-1].get('cikk_cim'):
        rows.pop()

    job_data = {
        'job_id': job_id,
        'status': 'pending',
        'started_at': datetime.now().isoformat(),
        'model': model,
        'total_rows': len(rows),
        'completed_rows': 0,
        'rows': rows
    }
    
    save_job(job_data)
    job_events[job_id] = Queue()

    thread = threading.Thread(target=generation_worker, args=(job_id,))
    thread.daemon = True
    thread.start()

    return jsonify({
        'success': True,
        'job_id': job_id
    })

# ==========================================
# ÚTVONALAK – JOB KEZELÉS
# ==========================================
@app.route('/jobs', methods=['GET'])
def list_jobs():
    jobs = get_all_jobs()
    # Csak a metaadatokat küldjük vissza, a nagy sor adatokat nem
    summary = []
    for j in jobs:
        summary.append({
            'job_id': j['job_id'],
            'status': j['status'],
            'started_at': j['started_at'],
            'total_rows': j['total_rows'],
            'completed_rows': j['completed_rows'],
            'download_url': j.get('download_url'),
            'zip_url': j.get('zip_url')
        })
    return jsonify({'jobs': summary})

@app.route('/jobs/<job_id>', methods=['GET'])
def get_job(job_id):
    job = load_job(job_id)
    if not job:
        return jsonify({'error': 'Job nem található'}), 404
    return jsonify(job)

@app.route('/jobs/<job_id>/resume', methods=['POST'])
def resume_job(job_id):
    if generation_lock.locked():
        return jsonify({'error': 'Már fut egy generálás. Kérlek várj.'}), 409
        
    job = load_job(job_id)
    if not job:
        return jsonify({'error': 'Job nem található'}), 404
        
    if job['status'] == 'done':
        return jsonify({'error': 'Ez a job már befejeződött'}), 400
        
    job['status'] = 'pending'
    save_job(job)
    
    if job_id not in job_events:
        job_events[job_id] = Queue()
        
    thread = threading.Thread(target=generation_worker, args=(job_id,))
    thread.daemon = True
    thread.start()
    
    return jsonify({
        'success': True,
        'job_id': job_id,
        'message': 'Job folytatása elindítva'
    })

@app.route('/jobs/<job_id>', methods=['DELETE'])
def delete_job(job_id):
    path = get_job_path(job_id)
    if os.path.exists(path):
        os.remove(path)
        return jsonify({'success': True})
    return jsonify({'error': 'Job nem található'}), 404

@app.route('/stream/<job_id>')
def stream(job_id):
    job = load_job(job_id)
    if not job:
        return jsonify({'error': 'Nem található ilyen folyamat'}), 404

    def event_stream():
        # Ha a job már kész, küldjük el az összes szükséges eseményt egyszerre
        if job['status'] == 'done':
            # Küldjük el a sorok végső állapotát
            for idx, row in enumerate(job['rows']):
                yield f"data: {json.dumps({'type': 'row_update', 'row_index': idx, 'status': row.get('status'), 'message': row.get('message', '')})}\n\n"
            
            # Küldjük el a progress-t
            yield f"data: {json.dumps({'type': 'progress', 'completed': job['completed_rows'], 'total': job['total_rows']})}\n\n"
            
            # Küldjük el a complete eseményt
            yield f"data: {json.dumps({'type': 'complete', 'download_url': job.get('download_url'), 'zip_url': job.get('zip_url')})}\n\n"
            return

        # Ha a job fut, csatlakozunk a sorhoz
        if job_id not in job_events:
            job_events[job_id] = Queue()
            
        queue = job_events[job_id]

        while True:
            event = queue.get()
            if event is None:
                break
            yield f"data: {json.dumps(event)}\n\n"

    return Response(event_stream(), mimetype="text/event-stream")

@app.route('/download/<filename>')
def download_file(filename):
    filepath = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
    if os.path.exists(filepath):
        # TTL ellenőrzés (1 óra)
        file_time = os.path.getmtime(filepath)
        if time.time() - file_time > 3600:
            # A fájl lejárt, de nem töröljük automatikusan letöltéskor,
            # inkább csak jelezzük a felhasználónak (opcionális extra biztonság)
            pass
            
        return send_file(filepath, as_attachment=True)
    return "Fájl nem található vagy már törölték a szerverről (1 órás limit).", 404

# ==========================================
# ÚTVONALAK – PIPELINE & TONE GUIDE
# ==========================================
@app.route('/pipeline', methods=['GET'])
def get_pipeline():
    data = load_pipeline_data()
    return jsonify(data)

@app.route('/pipeline', methods=['POST'])
def update_pipeline():
    data = request.json
    steps = data.get('steps', [])
    if not steps:
        return jsonify({'error': 'A pipeline nem lehet üres'}), 400

    save_pipeline_data(steps)
    return jsonify({'success': True, 'message': 'Pipeline sikeresen mentve'})

@app.route('/pipeline/versions', methods=['GET'])
def get_pipeline_versions():
    data = load_pipeline_data()
    return jsonify({
        'versions': data.get('versions', [])
    })

@app.route('/pipeline/restore/<int:version_number>', methods=['POST'])
def restore_pipeline_version(version_number):
    data = load_pipeline_data()
    versions = data.get('versions', [])

    target = next((v for v in versions if v['version'] == version_number), None)
    if not target:
        return jsonify({'error': f'Nem található {version_number}. verzió'}), 404

    save_pipeline_data(target['steps'])
    return jsonify({'success': True, 'message': f'{version_number}. verzió visszaállítva'})

@app.route('/prompts/tone_guide', methods=['GET'])
def get_tone_guide_route():
    return jsonify({'text': load_tone_guide()})

@app.route('/prompts/tone_guide', methods=['POST'])
def update_tone_guide():
    data = request.json
    new_text = data.get('text', '').strip()
    if not new_text:
        return jsonify({'error': 'A tone guide nem lehet üres'}), 400

    save_tone_guide(new_text)
    return jsonify({'success': True, 'message': 'Tone guide mentve'})

@app.route('/variables', methods=['GET'])
def get_variables():
    variables = {
        "Excel változók": {
            "{ceg_url}": "A cég honlapjának URL-je",
            "{cikk_cim}": "A cikk H1 címe",
            "{link_N_kulcsszo}": "Belső link kulcsszava (N = 1-5)",
            "{link_N_url}": "Belső link URL-je (N = 1-5)",
            "{korabbi_cikk_url_N}": "Korábbi cikk URL-je (N = 1-2)",
            "{megjegyzes}": "Egyéb instrukció az Excelből"
        },
        "Generált változók": {
            "{linkek_felsorolasa}": "Automatikus lista a megadott linkekből Markdown formátumban",
            "{opcionalis_korabbi_cikkek_blokk}": "Korábbi cikkek hivatkozási instrukciója (ha van)",
            "{megjegyzes_blokk}": "Megjegyzés instrukció (ha van)",
            "{elvart_linkek}": "A megadott kulcsszavak listája ellenőrzéshez"
        },
        "Rendszer változók": {
            "{tone_guide}": "A Tone Guide teljes szövege",
            "{aktualis_cikk}": "Az előző lépés kimenete (ez lesz a végső cikk az utolsó lépés után)",
            "{elozo_kimenet}": "A közvetlenül előző lépés teljes GPT válasza",
            "{lepes_N_kimenet}": "Az N. sorszámú (vagy azonosítójú) lépés GPT válasza"
        }
    }
    return jsonify(variables)

# ==========================================
# INDÍTÁS
# ==========================================
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
